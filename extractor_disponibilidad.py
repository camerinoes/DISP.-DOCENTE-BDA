"""
Extractor de disponibilidad horaria docente
BDA Analytics — Universidad de Lima

Uso:
    python extractor_disponibilidad.py --carpeta ./word_files --salida disponibilidad.xlsx

Si no se pasa --carpeta, busca archivos .docx en el directorio actual.
"""

import os
import re
import argparse
from pathlib import Path
from docx import Document
import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter


DIAS = ["LUNES", "MARTES", "MIÉRCOLES", "JUEVES", "VIERNES", "SÁBADO"]
HORAS = ["7-8","8-9","9-10","10-11","11-12","12-13","13-14",
         "14-15","15-16","16-17","17-18","18-19","19-20","20-21","21-22"]


def extraer_texto_campo(doc, etiqueta):
    """Busca una etiqueta en el texto del documento y devuelve el valor que le sigue."""
    texto_completo = " ".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texto_completo += " " + cell.text
    idx = texto_completo.upper().find(etiqueta.upper())
    if idx == -1:
        return ""
    fragmento = texto_completo[idx + len(etiqueta):idx + len(etiqueta) + 200]
    fragmento = fragmento.strip().split("\n")[0].strip()
    # Limpiar etiquetas siguientes comunes
    for siguiente in ["TELÉFONO", "CELULAR", "CORREO", "DATOS", "PERFIL", "MAESTRÍA",
                      "DOCTORADO", "EMPRESA", "CARGO", "AGREGUE"]:
        pos = fragmento.upper().find(siguiente)
        if pos > 0:
            fragmento = fragmento[:pos]
    return fragmento.strip().rstrip(":").strip()


def extraer_tabla_disponibilidad(doc):
    """Extrae la tabla de disponibilidad horaria."""
    disponibilidad = {dia: [] for dia in DIAS}

    for table in doc.tables:
        # Buscar la tabla de disponibilidad: primera fila debe tener HORAS y los días
        primera_fila = [c.text.strip().upper() for c in table.rows[0].cells]
        if "LUNES" in primera_fila and "HORAS" in primera_fila:
            idx_dia = {dia: primera_fila.index(dia) for dia in DIAS if dia in primera_fila}
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if not cells:
                    continue
                hora = cells[0].strip()
                if not hora or hora not in HORAS:
                    continue
                for dia, col_idx in idx_dia.items():
                    if col_idx < len(cells) and cells[col_idx].upper() in ("X", "✓", "✗", "x"):
                        disponibilidad[dia].append(hora)
            break
    return disponibilidad


def extraer_cursos_tabla(doc, etiqueta_busqueda):
    """Extrae cursos de una tabla que contenga la etiqueta buscada."""
    cursos = []
    for table in doc.tables:
        encabezados = [c.text.strip().upper() for c in table.rows[0].cells]
        if etiqueta_busqueda.upper() in " ".join(encabezados):
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if cells and cells[0] and cells[0].upper() not in ("ASIGNATURA", ""):
                    cursos.append(cells[0])
    return [c for c in cursos if c]


def extraer_jps(doc):
    """Extrae sugerencias de jefes de práctica."""
    jps = {}
    for table in doc.tables:
        encabezados = [c.text.strip().upper() for c in table.rows[0].cells]
        if "JEFE DE PRÁCTICAS" in " ".join(encabezados) or "JEFE DE PRACTICAS" in " ".join(encabezados):
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2 and cells[0] and cells[0].upper() not in ("ASIGNATURA", ""):
                    jps[cells[0]] = cells[1] if len(cells) > 1 else ""
    return jps


def disponibilidad_a_texto(disp):
    """Convierte el dict de disponibilidad a texto legible."""
    partes = []
    for dia in DIAS:
        horas = disp.get(dia, [])
        if horas:
            partes.append(f"{dia}: {', '.join(horas)}")
    return " | ".join(partes) if partes else "Sin disponibilidad marcada"


def disponibilidad_a_grilla(disp):
    """Devuelve dict {dia: set_de_horas} para comparación."""
    return {dia: set(horas) for dia, horas in disp.items()}


def procesar_archivo(ruta):
    """Procesa un archivo .docx y devuelve un dict con todos los datos."""
    try:
        doc = Document(ruta)
    except Exception as e:
        return {"archivo": str(ruta), "error": str(e)}

    nombre = extraer_texto_campo(doc, "APELLIDOS Y NOMBRES:")
    if not nombre:
        nombre = extraer_texto_campo(doc, "NOMBRES:")
    codigo = extraer_texto_campo(doc, "CÓDIGO:")
    correo = extraer_texto_campo(doc, "CORREO ELECTRÓNICO PERSONAL:")
    celular = extraer_texto_campo(doc, "CELULAR:")

    disponibilidad = extraer_tabla_disponibilidad(doc)
    cursos_asignados = extraer_cursos_tabla(doc, "NIVEL")
    cursos_adicionales = extraer_cursos_tabla(doc, "OTRAS ASIGNATURAS")
    jps = extraer_jps(doc)

    # Inglés y otras instituciones
    texto_doc = " ".join(p.text for p in doc.paragraphs)
    ingles = ""
    otras_instituciones = ""

    idx_ing = texto_doc.upper().find("INGLÉS")
    if idx_ing != -1:
        fragmento = texto_doc[idx_ing + 7:idx_ing + 80].strip()
        for stop in ["INDIQUE", "DICTA EN OTRAS", "INSTITUC", "AGRE"]:
            pos = fragmento.upper().find(stop)
            if pos > 0:
                fragmento = fragmento[:pos]
        ingles = fragmento.strip().rstrip(":,").strip()

    idx_otras = texto_doc.upper().find("OTRAS INSTITUCIONES")
    if idx_otras != -1:
        fragmento2 = texto_doc[idx_otras + 20:idx_otras + 150].strip()
        otras_instituciones = fragmento2.split("\n")[0].strip().rstrip(":").strip()

    return {
        "archivo": Path(ruta).name,
        "nombre": nombre,
        "codigo": codigo,
        "correo": correo,
        "celular": celular,
        "disponibilidad": disponibilidad,
        "disponibilidad_texto": disponibilidad_a_texto(disponibilidad),
        "cursos_asignados": cursos_asignados,
        "cursos_adicionales": cursos_adicionales,
        "jps": jps,
        "ingles": ingles,
        "otras_instituciones": otras_instituciones,
        "error": None,
    }


def exportar_excel(registros, ruta_salida):
    wb = openpyxl.Workbook()

    # ─── Hoja 1: Resumen por docente ───────────────────────────────────────
    ws = wb.active
    ws.title = "Disponibilidad"

    color_header = PatternFill("solid", fgColor="1F3864")
    color_disp   = PatternFill("solid", fgColor="C6EFCE")   # verde claro
    color_nodisp = PatternFill("solid", fgColor="FFFFFF")
    color_error  = PatternFill("solid", fgColor="FFD7D7")
    font_header  = Font(color="FFFFFF", bold=True, size=10)
    font_normal  = Font(size=9)
    centro       = Alignment(horizontal="center", vertical="center", wrap_text=True)

    # Cabecera
    cols_fijas = ["NOMBRE", "CÓDIGO", "CORREO", "CURSOS ASIGNADOS",
                  "CURSOS ADICIONALES", "JPs SUGERIDOS", "PUEDE EN INGLÉS"]
    cols_disp = [f"{dia}\n{h}" for dia in DIAS for h in HORAS]
    cabecera = cols_fijas + cols_disp
    ws.append(cabecera)

    for col_idx, titulo in enumerate(cabecera, 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.fill = color_header
        cell.font = font_header
        cell.alignment = centro

    # Filas de datos
    for r in registros:
        if r.get("error"):
            fila = [r["archivo"], "", "", f"ERROR: {r['error']}"] + [""] * (len(cabecera) - 4)
            ws.append(fila)
            for col_idx in range(1, len(cabecera) + 1):
                ws.cell(row=ws.max_row, column=col_idx).fill = color_error
            continue

        cursos_asig = "; ".join(r["cursos_asignados"])
        cursos_adic = "; ".join(r["cursos_adicionales"])
        jps_txt = "; ".join(f"{cur}: {jp}" for cur, jp in r["jps"].items())

        fila_base = [
            r["nombre"], r["codigo"], r["correo"],
            cursos_asig, cursos_adic, jps_txt, r["ingles"]
        ]

        disp_grilla = r["disponibilidad"]
        celdas_disp = []
        for dia in DIAS:
            horas_docente = set(disp_grilla.get(dia, []))
            for h in HORAS:
                celdas_disp.append("X" if h in horas_docente else "")

        ws.append(fila_base + celdas_disp)
        row_idx = ws.max_row
        for col_idx, val in enumerate(celdas_disp, len(cols_fijas) + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.alignment = centro
            cell.font = font_normal
            if val == "X":
                cell.fill = color_disp

        for col_idx in range(1, len(cols_fijas) + 1):
            ws.cell(row=row_idx, column=col_idx).font = font_normal
            ws.cell(row=row_idx, column=col_idx).alignment = Alignment(vertical="center", wrap_text=True)

    # Anchos de columna
    for col_idx in range(1, len(cols_fijas) + 1):
        ws.column_dimensions[get_column_letter(col_idx)].width = 22
    for col_idx in range(len(cols_fijas) + 1, len(cabecera) + 1):
        ws.column_dimensions[get_column_letter(col_idx)].width = 5
    ws.row_dimensions[1].height = 40

    # ─── Hoja 2: Grilla visual de disponibilidad ───────────────────────────
    ws2 = wb.create_sheet("Grilla por hora")
    color_h2 = PatternFill("solid", fgColor="1F3864")

    docentes_validos = [r for r in registros if not r.get("error")]
    nombres = [r["nombre"] or r["archivo"] for r in docentes_validos]

    # Cabecera: días y horas como filas, docentes como columnas
    ws2.cell(1, 1, "HORA").fill = color_h2
    ws2.cell(1, 1).font = font_header
    ws2.cell(1, 1).alignment = centro

    for col_idx, nombre in enumerate(nombres, 2):
        c = ws2.cell(1, col_idx, nombre)
        c.fill = color_h2
        c.font = font_header
        c.alignment = centro
        ws2.column_dimensions[get_column_letter(col_idx)].width = 20

    ws2.column_dimensions["A"].width = 18

    row = 2
    for dia in DIAS:
        for hora in HORAS:
            etiqueta = f"{dia} {hora}"
            ws2.cell(row, 1, etiqueta).alignment = centro
            ws2.cell(row, 1).font = Font(bold=True, size=9)
            for col_idx, r in enumerate(docentes_validos, 2):
                disponible = hora in set(r["disponibilidad"].get(dia, []))
                c = ws2.cell(row, col_idx, "✓" if disponible else "")
                c.alignment = centro
                c.font = Font(size=9)
                if disponible:
                    c.fill = color_disp
            row += 1

    # ─── Hoja 3: Índice de cursos ──────────────────────────────────────────
    ws3 = wb.create_sheet("Cursos por docente")
    ws3.append(["NOMBRE", "CÓDIGO", "CURSO ASIGNADO", "NIVEL", "TIPO"])
    for col_idx in range(1, 6):
        ws3.cell(1, col_idx).fill = color_header
        ws3.cell(1, col_idx).font = font_header
        ws3.cell(1, col_idx).alignment = centro

    for r in docentes_validos:
        for curso in r["cursos_asignados"]:
            ws3.append([r["nombre"], r["codigo"], curso, "", "ASIGNADO"])
        for curso in r["cursos_adicionales"]:
            ws3.append([r["nombre"], r["codigo"], curso, "", "ADICIONAL"])

    for col in ["A","B","C","D","E"]:
        ws3.column_dimensions[col].width = 28

    wb.save(ruta_salida)
    print(f"Excel guardado en: {ruta_salida}")


def main():
    parser = argparse.ArgumentParser(description="Extractor de disponibilidad docente")
    parser.add_argument("--carpeta", default=".", help="Carpeta con archivos .docx")
    parser.add_argument("--salida", default="disponibilidad_consolidada.xlsx", help="Archivo Excel de salida")
    args = parser.parse_args()

    carpeta = Path(args.carpeta)
    archivos = list(carpeta.glob("*.docx")) + list(carpeta.glob("*.DOCX"))

    if not archivos:
        print(f"No se encontraron archivos .docx en: {carpeta}")
        return

    print(f"Procesando {len(archivos)} archivos...")
    registros = []
    for ruta in sorted(archivos):
        print(f"  → {ruta.name}")
        registros.append(procesar_archivo(ruta))

    exportar_excel(registros, args.salida)
    print(f"\nResumen:")
    print(f"  Procesados: {len(registros)}")
    print(f"  Con errores: {sum(1 for r in registros if r.get('error'))}")
    print(f"  Exitosos:    {sum(1 for r in registros if not r.get('error'))}")


if __name__ == "__main__":
    main()
